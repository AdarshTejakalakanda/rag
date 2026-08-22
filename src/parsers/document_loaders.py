"""Document loader abstractions supporting Markdown, TXT, PDF, and DOCX.

Conforms to Specification §4.
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional


class BaseDocumentLoader(ABC):
    """Abstract base class for local document loaders."""

    @abstractmethod
    def load(self, file_path: str or Path) -> str:
        """Loads and returns raw text content from the file."""
        pass


class TextDocumentLoader(BaseDocumentLoader):
    """Loads plain text files (.txt)."""

    def load(self, file_path: str or Path) -> str:
        path = Path(file_path)
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()


class MarkdownDocumentLoader(BaseDocumentLoader):
    """Loads Markdown business specification documents (.md, .markdown)."""

    def load(self, file_path: str or Path) -> str:
        path = Path(file_path)
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()


class PDFDocumentLoader(BaseDocumentLoader):
    """Loads PDF business requirement documents using pypdf."""

    def load(self, file_path: str or Path) -> str:
        path = Path(file_path)
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            pages_text = []
            for idx, page in enumerate(reader.pages):
                txt = page.extract_text() or ""
                if txt.strip():
                    pages_text.append(f"--- Page {idx + 1} ---\n{txt}")
            return "\n\n".join(pages_text)
        except Exception as e:
            raise RuntimeError(f"Failed loading PDF document {file_path}: {e}")


class DocxDocumentLoader(BaseDocumentLoader):
    """Loads Microsoft Word documents (.docx) using python-docx."""

    def load(self, file_path: str or Path) -> str:
        path = Path(file_path)
        try:
            import docx
            doc = docx.Document(str(path))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    full_text.append(" | ".join(row_text))
            return "\n\n".join(full_text)
        except Exception as e:
            raise RuntimeError(f"Failed loading DOCX document {file_path}: {e}")


class DocumentLoaderFactory:
    """Factory creating appropriate document loader based on file format."""

    _LOADERS = {
        ".txt": TextDocumentLoader,
        ".md": MarkdownDocumentLoader,
        ".markdown": MarkdownDocumentLoader,
        ".pdf": PDFDocumentLoader,
        ".docx": DocxDocumentLoader,
    }

    @classmethod
    def get_loader(cls, file_path: str or Path) -> BaseDocumentLoader:
        suffix = Path(file_path).suffix.lower()
        loader_cls = cls._LOADERS.get(suffix)
        if not loader_cls:
            # Fallback to Text loader
            return TextDocumentLoader()
        return loader_cls()

    @classmethod
    def load_file(cls, file_path: str or Path) -> str:
        loader = cls.get_loader(file_path)
        return loader.load(file_path)
