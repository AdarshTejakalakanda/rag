"""Tests for Multi-Format Document Loaders conforming to Specification §4."""

import pytest
from pathlib import Path
from src.parsers.document_loaders import (
    DocumentLoaderFactory,
    TextDocumentLoader,
    MarkdownDocumentLoader,
    DocxDocumentLoader,
)
from src.parsers.requirement_parser import RequirementParser


def test_document_loaders(tmp_path):
    # 1. Text Document
    txt_file = tmp_path / "spec.txt"
    txt_file.write_text("REQ-01: User Login\nUser must login with valid password.", encoding="utf-8")
    txt_loader = DocumentLoaderFactory.get_loader(txt_file)
    assert isinstance(txt_loader, TextDocumentLoader)
    assert "REQ-01" in txt_loader.load(txt_file)

    # 2. Markdown Document
    md_file = tmp_path / "spec.md"
    md_file.write_text("# Feature\n## REQ-02: Checkout\nCart total must calculate tax.", encoding="utf-8")
    md_loader = DocumentLoaderFactory.get_loader(md_file)
    assert isinstance(md_loader, MarkdownDocumentLoader)
    assert "REQ-02" in md_loader.load(md_file)

    # 3. DOCX Document (create simple docx)
    import docx
    docx_file = tmp_path / "spec.docx"
    doc = docx.Document()
    doc.add_heading("E-Commerce BRD", level=1)
    doc.add_paragraph("REQ-03: Member 360 Alerts\nUsers can create member alerts.")
    doc.save(str(docx_file))

    docx_loader = DocumentLoaderFactory.get_loader(docx_file)
    assert isinstance(docx_loader, DocxDocumentLoader)
    loaded_docx = docx_loader.load(docx_file)
    assert "REQ-03" in loaded_docx
    assert "Member 360" in loaded_docx

    # 4. RequirementParser parsing docx
    reqs = RequirementParser.parse_file(docx_file)
    assert len(reqs) >= 1
    assert any(r.req_id == "REQ-03" for r in reqs)
